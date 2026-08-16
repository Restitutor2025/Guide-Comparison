import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT

from guide_comparison.diff.models import AlignedPair, DiffStatus, TextBlock
from guide_comparison.reporting import build_report_rows, export_word_report


class ReportingTests(unittest.TestCase):
    def test_report_rows_are_editable_copies_and_skip_unchanged_pairs(self):
        old = TextBlock("최대 3개월", page=12)
        new = TextBlock("최대 2개월", page=13)
        pairs = [
            AlignedPair(old, new, DiffStatus.MODIFIED),
            AlignedPair(TextBlock("동일", page=14), TextBlock("동일", page=14), DiffStatus.UNCHANGED),
        ]

        rows = build_report_rows(pairs)

        self.assertEqual(len(rows), 1)
        self.assertEqual((rows[0].status, rows[0].old_page, rows[0].new_page), ("수정", "12", "13"))
        rows[0].old_text = "보고용 수정 문구"
        self.assertEqual(old.text, "최대 3개월")

    def test_word_report_contains_only_included_review_rows(self):
        pairs = [
            AlignedPair(TextBlock("기존 기준", page=1), TextBlock("개정 기준", page=2), DiffStatus.MODIFIED),
            AlignedPair(None, TextBlock("신규 기준", page=3), DiffStatus.ADDED),
        ]
        rows = build_report_rows(pairs)
        rows[0].review_note = "기간 단축 여부 검토"
        rows[1].included = False

        with tempfile.TemporaryDirectory() as folder:
            path = Path(folder) / "report.docx"
            export_word_report(path, rows, "old.pdf", "new.pdf")
            document = Document(path)

        self.assertEqual(document.sections[0].orientation, WD_ORIENT.LANDSCAPE)
        self.assertEqual(len(document.tables), 1)
        self.assertEqual(len(document.tables[0].rows), 2)
        table_text = "\n".join(cell.text for row in document.tables[0].rows for cell in row.cells)
        self.assertIn("기존 기준", table_text)
        self.assertIn("개정 기준", table_text)
        self.assertIn("기간 단축 여부 검토", table_text)
        self.assertNotIn("신규 기준", table_text)


if __name__ == "__main__":
    unittest.main()
