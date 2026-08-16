import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from guide_comparison.diff.models import TableBlock, TextBlock
from guide_comparison.parsers import DocumentParseError, parse_document
from guide_comparison.parsers.docx_parser import _extract_docx_blocks
from guide_comparison.parsers.pdf_parser import _assign_structure_paths, _merge_page_blocks


class ParserTests(unittest.TestCase):
    def test_unsupported_extension(self):
        with tempfile.TemporaryDirectory() as folder:
            path = Path(folder) / "data.txt"; path.write_text("hello")
            with self.assertRaises(DocumentParseError):
                parse_document(path)

    def test_docx_preserves_paragraph_table_order(self):
        try:
            from docx import Document
        except ImportError:
            self.skipTest("python-docx unavailable")
        with tempfile.TemporaryDirectory() as folder:
            path = Path(folder) / "order.docx"
            doc = Document(); doc.add_paragraph("Before"); table = doc.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "A"; table.cell(0, 1).text = "B"; doc.add_paragraph("After"); doc.save(path)
            source_blocks = _extract_docx_blocks(path)
            self.assertEqual([type(b) for b in source_blocks], [TextBlock, TableBlock, TextBlock])
            import pymupdf as fitz
            rendered_path = Path(folder) / "rendered.pdf"
            rendered = fitz.open(); page = rendered.new_page(); page.insert_text((72, 72), "Before A B After")
            rendered.save(rendered_path); rendered.close()
            with patch("guide_comparison.parsers.docx_parser._render_docx_to_pdf", return_value=rendered_path):
                parsed = parse_document(path)
            self.assertEqual(parsed.page_count, 1)
            self.assertTrue(parsed.render_path and parsed.render_path.is_file())
            self.assertIn("Before", " ".join(block.text for block in parsed.blocks))

    def test_pdf_page_count_and_text(self):
        try:
            import pymupdf as fitz
        except ImportError:
            self.skipTest("PyMuPDF unavailable")
        with tempfile.TemporaryDirectory() as folder:
            path = Path(folder) / "sample.pdf"
            doc = fitz.open(); page = doc.new_page(); page.insert_text((72, 72), "Guide paragraph"); doc.save(path); doc.close()
            parsed = parse_document(path)
            self.assertEqual(parsed.page_count, 1)
            self.assertIn("Guide paragraph", [b.text for b in parsed.blocks])
            self.assertEqual(parsed.render_path, path)
            self.assertTrue(next(b for b in parsed.blocks if b.text == "Guide paragraph").rects)

    def test_pdf_merges_wrapped_lines_but_keeps_separate_paragraphs(self):
        try:
            import pymupdf as fitz
        except ImportError:
            self.skipTest("PyMuPDF unavailable")
        with tempfile.TemporaryDirectory() as folder:
            path = Path(folder) / "wrapped.pdf"
            doc = fitz.open(); page = doc.new_page()
            page.insert_text((72, 72), "First paragraph continues")
            page.insert_text((72, 86), "and ends here.")
            page.insert_text((72, 125), "Separate paragraph.")
            doc.save(path); doc.close()
            parsed = parse_document(path)
            self.assertEqual([block.text for block in parsed.blocks], [
                "First paragraph continues and ends here.",
                "Separate paragraph.",
            ])
            self.assertEqual(len(parsed.blocks[0].rects), 2)
            self.assertTrue(parsed.blocks[0].chars)

    def test_list_continuations_merge_but_new_structural_markers_split(self):
        blocks = [
            TextBlock("② Existing requirement", "list", 1, [(72, 60, 250, 75)]),
            TextBlock("continues on the next line", "paragraph", 1, [(96, 75, 270, 90)]),
            TextBlock("※ New independent note", "list", 1, [(72, 90, 260, 105)]),
        ]
        merged = _merge_page_blocks(blocks, 600, [])
        self.assertEqual([block.text for block in merged], [
            "② Existing requirement continues on the next line",
            "※ New independent note",
        ])

    def test_pdf_blocks_receive_heading_and_nested_item_structure_paths(self):
        blocks = [
            TextBlock("3 평가 개요", "heading", 1, [(50, 20, 200, 40)]),
            TextBlock("① 적용범위", "list", 1, [(70, 60, 200, 75)]),
            TextBlock("○ 설명 문구", "list", 1, [(85, 85, 260, 100)]),
            TextBlock("< 평가가 필요한 변경사항 >", "heading", 1, [(100, 120, 400, 140)]),
            TextBlock("② 패키지 구성품이 추가된 경우", "list", 1, [(110, 155, 400, 170)]),
            TextBlock("③ GMP 평가를 받지 않은 경우", "list", 1, [(110, 180, 400, 195)]),
            TextBlock("[예시]", "paragraph", 1, [(80, 220, 130, 235)]),
        ]

        _assign_structure_paths(blocks)

        self.assertEqual(blocks[0].structure_path, ("heading:number:3",))
        self.assertEqual(blocks[1].structure_path, ("heading:number:3", "item:①"))
        self.assertEqual(blocks[2].structure_path, ("heading:number:3", "item:①", "item:○"))
        self.assertEqual(
            blocks[4].structure_path,
            ("heading:number:3", "heading:box:평가가필요한변경사항", "item:②"),
        )
        self.assertEqual(
            blocks[5].structure_path,
            ("heading:number:3", "heading:box:평가가필요한변경사항", "item:③"),
        )
        self.assertEqual(
            blocks[6].structure_path,
            ("heading:number:3", "heading:box:평가가필요한변경사항", "item:[예시]"),
        )

    def test_corrupt_files_are_wrapped(self):
        with tempfile.TemporaryDirectory() as folder:
            for suffix in (".pdf", ".docx"):
                path = Path(folder) / f"bad{suffix}"; path.write_bytes(b"not a document")
                with self.assertRaises(DocumentParseError):
                    parse_document(path)


if __name__ == "__main__":
    unittest.main()
