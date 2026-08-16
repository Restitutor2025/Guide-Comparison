from __future__ import annotations

from collections import Counter
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from guide_comparison.diff.models import AlignedPair, DiffStatus, TableBlock, TextBlock


STATUS_LABELS = {
    DiffStatus.MODIFIED: "수정",
    DiffStatus.ADDED: "추가",
    DiffStatus.REMOVED: "삭제",
}


@dataclass(slots=True)
class ReportRow:
    included: bool
    number: int
    status: str
    old_page: str
    new_page: str
    old_text: str
    new_text: str
    review_note: str = ""


def _block_text(block) -> str:
    if isinstance(block, TextBlock):
        return block.text.strip()
    if isinstance(block, TableBlock):
        return "\n".join(" | ".join(cell.strip() for cell in row) for row in block.rows).strip()
    return ""


def build_report_rows(pairs: list[AlignedPair]) -> list[ReportRow]:
    """Copy detected changes into editable report rows without mutating source blocks."""
    rows: list[ReportRow] = []
    for pair in pairs:
        if pair.status == DiffStatus.UNCHANGED:
            continue
        old_page = getattr(pair.old_block, "page", None)
        new_page = getattr(pair.new_block, "page", None)
        rows.append(ReportRow(
            included=True,
            number=len(rows) + 1,
            status=STATUS_LABELS.get(pair.status, "확인 필요"),
            old_page=str(old_page) if old_page is not None else "-",
            new_page=str(new_page) if new_page is not None else "-",
            old_text=_block_text(pair.old_block) or "해당 없음",
            new_text=_block_text(pair.new_block) or "해당 없음",
        ))
    return rows


def _set_run_font(run, size: float, *, bold: bool = False, color: str = "20262E") -> None:
    run.font.name = "맑은 고딕"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Malgun Gothic")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Malgun Gothic")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell, *, top: int = 80, start: int = 100, bottom: int = 80, end: int = 100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, (cell, width) in enumerate(zip(row.cells, widths_dxa)):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Cm(width / 1440 * 2.54)


def _write_cell(cell, text: str, *, bold: bool = False, centered: bool = False, size: float = 8.5) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.08
    _set_run_font(paragraph.add_run(text), size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_cell_margins(cell)


def export_word_report(
    destination: str | Path,
    rows: list[ReportRow],
    old_name: str,
    new_name: str,
) -> Path:
    """Create a new DOCX from reviewed rows; source documents are never opened for writing."""
    destination = Path(destination)
    selected = [row for row in rows if row.included]
    if not selected:
        raise ValueError("보고서에 포함할 변경점이 없습니다.")

    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.3)
    section.bottom_margin = Cm(1.3)
    section.left_margin = Cm(1.3)
    section.right_margin = Cm(1.3)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)

    normal = document.styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(9)

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    _set_run_font(title.add_run("문서 변경점 검토 결과"), 17, bold=True, color="17365D")

    metadata = document.add_paragraph()
    metadata.paragraph_format.space_after = Pt(7)
    _set_run_font(metadata.add_run(
        f"기존 문서: {old_name}\n개정 문서: {new_name}\n작성일: {datetime.now():%Y-%m-%d %H:%M}"
    ), 9, color="536170")

    counts = Counter(row.status for row in selected)
    summary = document.add_paragraph()
    summary.paragraph_format.space_after = Pt(7)
    summary_text = "  |  ".join([
        f"전체 {len(selected)}건",
        f"수정 {counts.get('수정', 0)}건",
        f"추가 {counts.get('추가', 0)}건",
        f"삭제 {counts.get('삭제', 0)}건",
        f"기타 {len(selected) - counts.get('수정', 0) - counts.get('추가', 0) - counts.get('삭제', 0)}건",
    ])
    _set_run_font(summary.add_run(summary_text), 9, bold=True, color="1F4D78")

    notice = document.add_paragraph()
    notice.paragraph_format.space_after = Pt(8)
    _set_run_font(
        notice.add_run("※ 사용자가 검토 화면에서 편집·선택한 결과이며 원본 파일은 변경되지 않았습니다."),
        8.5,
        color="6A737D",
    )

    headers = ["번호", "구분", "문서 위치", "기존 내용", "변경 내용", "변경 요약·검토 의견"]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_row = table.rows[0]
    _set_repeat_table_header(header_row)
    for cell, header in zip(header_row.cells, headers):
        _write_cell(cell, header, bold=True, centered=True, size=8.5)
        _set_cell_shading(cell, "DCE6F1")

    status_fills = {"수정": "FFF2CC", "추가": "E2F0D9", "삭제": "FCE4D6", "확인 필요": "DDEBF7"}
    for output_number, row in enumerate(selected, 1):
        cells = table.add_row().cells
        old_location = f"기존 p.{row.old_page}" if row.old_page not in {"", "-"} else "기존 -"
        new_location = f"개정 p.{row.new_page}" if row.new_page not in {"", "-"} else "개정 -"
        values = [
            str(output_number),
            row.status,
            f"{old_location}\n{new_location}",
            row.old_text,
            row.new_text,
            row.review_note,
        ]
        for index, (cell, value) in enumerate(zip(cells, values)):
            _write_cell(cell, value, centered=index in {0, 1, 2})
        _set_cell_shading(cells[1], status_fills.get(row.status, "DDEBF7"))

    # A4 landscape usable width: 27.1 cm / 15,364 DXA.
    _set_table_geometry(table, [600, 900, 1500, 4300, 4300, 3764])

    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)
    return destination
